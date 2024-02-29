import requests
from bs4 import BeautifulSoup
import re
from docx import Document
import os
from pypinyin import lazy_pinyin
import chardet

session = requests.Session()  # 使用 Session 对象
def get_html(url):
    try:
        response = session.get(url)  # 使用 session 发起请求
        response.raise_for_status()

        # 使用 chardet 检测编码
        detected_encoding = chardet.detect(response.content)['encoding']
        if detected_encoding:
            response.encoding = detected_encoding
        else:
            response.encoding = 'gb2312'  # 如果检测失败，默认使用 GB2312

        return response.text  # 使用正确的编码解码文本
    except requests.RequestException as e:
        return f"Error: {e}"

def get_final_number_and_filename(url,get_html):
    html_content = get_html(url)
    soup = BeautifulSoup(html_content, 'lxml')
    nav_element = soup.select_one('nav.pagination.is-small.is-centered')
    a_tags = nav_element.find_all('a')
    # 获取倒数第二个<a>标签的内容
    finalNumber = a_tags[-2].get_text()
    title_tag = soup.find('title')
    filename = title_tag.get_text().strip()
    filename=filename+".docx"
    return finalNumber,filename

def get_answer(url, get_html, get_final_number_and_filename):
    url_list = []
    final_number,filename = get_final_number_and_filename(url, get_html)
    filename=str(filename)
    all_p_content = ""  # 定义中间变量来累积所有 <p> 标签的内容

    for i in range(1, int(final_number) + 1):
        pattern = r"(.*/)(\d{6})"
        match = re.search(pattern, url)
        if match:
            extracted_url = match.group(1) + match.group(2)
            new_url = extracted_url + (".html" if i == 1 else f"_{i}.html")
            url_list.append(new_url)
        else:
            print("没有找到匹配的URL部分")

    for url in url_list:
        html_content = get_html(url)
        soup = BeautifulSoup(html_content, 'lxml')

        answer_content = soup.select_one('div.content.is-normal')
        p_tags = answer_content.find_all('p')

        for p in p_tags:
            all_p_content += p.get_text() + "\n"  # 追加 <p> 标签内容和换行符

    return all_p_content,filename


def sort_questions(content):
    pattern = re.compile(r"\【.*?\】.*?我的答案：[^\n]*", re.DOTALL | re.MULTILINE)
    questions = pattern.findall(content)

    questions_with_pinyin = []

    for question in questions:
        # 提取题干部分
        title_start = question.find('】') + 1
        title_end = question.find('\n', title_start)
        title = question[title_start:title_end]

        # 提取题干的首字拼音并转为小写
        pinyin = lazy_pinyin(title[0])[0].lower() if title else ''

        # 将题目和首字的拼音添加到列表中
        questions_with_pinyin.append((question, pinyin))

    # 根据拼音排序
    sorted_questions = sorted(questions_with_pinyin, key=lambda x: x[1])
    return sorted_questions


def save_to_word(sorted_questions, filename):
    doc = Document()
    current_initial = ''
    for question, pinyin in sorted_questions:
        initial = pinyin[0].upper() if pinyin else '#'

        if initial != current_initial:
            doc.add_heading(f"{initial}" + "-" * 18 + f"{initial}", level=1)
            current_initial = initial

        doc.add_paragraph(question)
    # 保存到桌面
    desktop = os.path.join(os.path.join(os.environ['USERPROFILE']), 'Desktop')
    doc.save(os.path.join(desktop, filename))

if __name__ == '__main__':
    url=input("请输入url:")
    all_p_content,filename=get_answer(url,get_html,get_final_number_and_filename)
    sorted_questions=sort_questions(all_p_content)
    save_to_word(sorted_questions, filename)
    print("文档已保存在桌面")