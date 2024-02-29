from docx import Document
from pypinyin import lazy_pinyin
import re
import os
def get_docx_files(directory):
    """获取指定目录下所有的docx文件的路径"""
    docx_files = []
    for file in os.listdir(directory):
        if file.endswith(".docx"):
            docx_files.append(os.path.join(directory, file))
    return docx_files

def read_docx(file_paths):
    """读取多个文档的内容并合并"""
    full_text = []
    for file_path in file_paths:
        doc = Document(file_path)
        for para in doc.paragraphs:
            full_text.append(para.text)
    return '\n'.join(full_text)

def sort_questions(content):
    pattern = r'(\d+\n[\s\S]*?我的答案：[A-Za-z]+\n)'
    questions = re.findall(pattern, content, re.DOTALL)

    questions_with_pinyin = []

    for question in questions:
        # 提取题干部分
        # 找到第一个和第二个换行符的位置
        title_start = question.find('\n') + 1
        title_end = question.find('\n', title_start)
        title = question[title_start:title_end]

        # 提取题干的首字拼音并转为小写
        pinyin = lazy_pinyin(title[0])[0].lower() if title else ''

        # 将题目和首字的拼音添加到列表中
        questions_with_pinyin.append((question, pinyin))

    # 根据拼音排序
    sorted_questions = sorted(questions_with_pinyin, key=lambda x: x[1])
    return sorted_questions



def save_to_word(sorted_questions):
    doc = Document()
    current_initial = ''
    for question, pinyin in sorted_questions:
        initial = pinyin[0].upper() if pinyin else '#'

        if initial != current_initial:
            doc.add_heading(f"{initial}" + "-" * 18 + f"{initial}", level=1)
            current_initial = initial

        doc.add_paragraph(question)
    # 保存到桌面
    desktop = os.path.join(os.path.join(os.environ['USERPROFILE']), 'Desktop')
    doc.save(os.path.join(desktop, "《数学大观》期末考试答案_1.docx"))

if __name__ == '__main__':
    directory_path = 'C:/Users/24766/Desktop/新建文件夹'
    docx_files = get_docx_files(directory_path)

    doc_content = read_docx(docx_files)
    sort_questions=sort_questions(doc_content)
    save_to_word(sort_questions)